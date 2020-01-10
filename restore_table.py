import pickle

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt, Inches


# 设置表格行高
def set_row_height(table, begin, end, ratio):
    # 表格元素、起始行、终止行、行占原图比例
    try:
        for row in range(begin, end):
            row = table.rows[row]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(ratio*13000)))  # 13000为在纸张21.5*27.9cm, 上下距离25.4mm时页面设置最大值

            # trHeight.set(qn('w:val'), str(30))  # 强制最小 根据cell内容自适应
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
    except Exception as ex:
        print('set_row_height', ex)


# 设置表格列宽
def set_column_width(table, begin, end, width):
    try:
        for col in range(begin, end):
            # WPS设置列宽
            table.columns[col].width = width
            # office2016设置列宽
            rows = table.rows
            for row in range(len(rows)):
                table.cell(row, col).width = width
    except Exception as ex:
        print('set_column_width', ex)


def restore_table(doc, i, img):

    i = ['table', i[0]]
    table = doc.add_table(i[1][1], i[1][2], style='Table Grid')

    # 表格居中
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for j in i[1][0]:
        text = ''
        d = j[0]
        # 设置行高
        if d['col_begin'] == 0:
            aa = 1
            if d['row_end'] - d['row_begin'] != 0:
                aa = d['row_end'] - d['row_begin']
            row_size = j[1][2]/(img.height*aa)
            set_row_height(table, d['row_begin'], d['row_end'], row_size)

        # 设置列宽
        if d['row_begin'] == 0:
            bb = 1
            if d['col_end'] - d['col_begin'] != 0:
                bb = d['col_end'] - d['col_begin']
            col_size = 8 * (j[1][1] / bb) / img.width
            set_column_width(table, d['col_begin'], d['col_end'], Inches(col_size))

        # for col in range(d['col_begin'], d['col_end']):
        #     for row in range(d['row_begin'], d['row_end']):
        #         cell = table.cell(d['row_begin'], d['col_begin'])
        #         cell.merge(table.cell(row, col))
        #         if not cell.text:
        #             text = j[1][0]

        # 向单元格中添加值
        try:
            cell = table.cell(d['row_begin'], d['col_begin'])
            cell.merge(table.cell(d['row_end']-1, d['col_end']-1))
            if not cell.text:
                text = j[1][0]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(7)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing_rule = 0
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        except Exception as ex:
            print('cell_error', ex)
            # cell = table.cell(d['row_begin'], d['col_begin'])
            # if not cell.text:
            #     text = j[1][0]
    return doc


def restore_table_fast_ocr(doc, i, img):

    i = ['table', i[0]]
    table = doc.add_table(i[1][1], i[1][2], style='Table Grid')

    # 表格居中
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for j in i[1][0]:
        text = ''
        d = j[0]
        # 设置行高
        if d['col_begin'] == 0:
            aa = 1
            if d['row_end'] - d['row_begin'] != 0:
                aa = d['row_end'] - d['row_begin']
            row_size = j[1][2]/(img.height*aa)
            set_row_height(table, d['row_begin'], d['row_end'], row_size)

        # 设置列宽
        if d['row_begin'] == 0:
            bb = 1
            if d['col_end'] - d['col_begin'] != 0:
                bb = d['col_end'] - d['col_begin']
            col_size = 8 * (j[1][1] / bb) / img.width
            set_column_width(table, d['col_begin'], d['col_end'], Inches(col_size))

        # for col in range(d['col_begin'], d['col_end']):
        #     for row in range(d['row_begin'], d['row_end']):
        #         cell = table.cell(d['row_begin'], d['col_begin'])
        #         cell.merge(table.cell(row, col))
        #         if not cell.text:
        #             text = j[1][0]

        # 向单元格中添加值
        try:
            cell = table.cell(d['row_begin'], d['col_begin'])
            cell.merge(table.cell(d['row_end']-1, d['col_end']-1))
            if not cell.text:
                text = j[1][0]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(7)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing_rule = 0
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        except Exception as ex:
            print('cell_error', ex)
            # cell = table.cell(d['row_begin'], d['col_begin'])
            # if not cell.text:
            #     text = j[1][0]

    return doc


if __name__ == '__main__':
    img = Image.open(r'C:\Users\Admin\Desktop\texts.jpg')
    img.thumbnail((2500, 2500))
    doc = Document()
    data = pickle.load(open(r'F:\paragraph_restore\pdf\p1\76\76_103.pkl', 'rb'))
    data = [i for i in data if i[0] == 'table']
    for i in data:
        doc = restore_table(doc, i, img)
    doc.save('1.docx')
    pass
